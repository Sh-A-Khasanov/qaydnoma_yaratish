import tkinter as tk
from tkinter import ttk
import pandas as pd
import glob
import os
from docx import Document
from tkcalendar import DateEntry
from tkinter import messagebox

from word_yasash import create_word,delete_file


import requests
from io import StringIO

import pandas as pd
import requests
from io import StringIO


def get_uqituvchi_list_from_google_sheet():
    sheet_id = "1cGNF3MPX5agBNJmSJWfj1Tc0HVhIwTp338cRWzvkpgI"  # Yangi sheet_id
    sheet_name = "Xodimlar"
    url = f"https://docs.google.com/spreadsheets/d/{sheet_id}/gviz/tq?tqx=out:csv&sheet={sheet_name}"

    try:
        response = requests.get(url)
        response.raise_for_status()
        df = pd.read_csv(StringIO(response.text))

        # Ustun nomlari: Ismi (G), Familya (H), Otasining ismi (I)
        if {'Ismi', 'Familiya', "Otasining ismi"}.issubset(df.columns):
            full_names = df[['Familiya', 'Ismi', 'Otasining ismi']].fillna('').astype(str).apply(
                lambda x: f"{x['Familiya']} {x['Ismi']} {x['Otasining ismi']}".strip(), axis=1)
            return full_names.tolist()
        else:
            print("❌ Kutilgan ustun nomlari topilmadi.")
            print("Topilgan ustunlar:", df.columns.tolist())
            return []
    except Exception as e:
        print("❌ Google Sheet dan o‘qishda xatolik:", e)
        return []

def get_fanlar_from_google_sheet():
    sheet_id = "1cGNF3MPX5agBNJmSJWfj1Tc0HVhIwTp338cRWzvkpgI"  # Yangi sheet_id
    sheet_name = "Fanlar ro'yxati"
    url = f"https://docs.google.com/spreadsheets/d/{sheet_id}/gviz/tq?tqx=out:csv&sheet={sheet_name}"

    try:
        response = requests.get(url)
        response.raise_for_status()
        df = pd.read_csv(StringIO(response.text))

        if 'Nomi' in df.columns:
            fanlar = df['Nomi'].dropna().astype(str).tolist()
            return fanlar
        else:
            print("❌ 'Nomi' ustuni topilmadi.")
            print("Topilgan ustunlar:", df.columns.tolist())
            return []
    except Exception as e:
        print("❌ Fanlar listidan o‘qishda xatolik:", e)
        return []


# Excel faylini avtomatik tanlash
def get_latest_excel_file(folder="data"):
    list_of_files = glob.glob(os.path.join(folder, "Talabalar-*.xlsx"))
    if not list_of_files:
        return None
    latest_file = max(list_of_files, key=os.path.getctime)
    return latest_file


# Excel fayllarni yuklash
talabalar = get_latest_excel_file("data")
if not talabalar:
    print("❌ Talabalar fayli topilmadi!")
    exit()



# Guruh va fakultet nomlarini olish
def get_groups_and_faculties_from_excel(file_path):
    try:
        df = pd.read_excel(file_path)
        try:
            groups = df['O'].dropna().unique()
        except:
            groups = df.iloc[:, 14].dropna().unique()
        try:
            faculties = df['N'].dropna().unique()
        except:
            faculties = df.iloc[:, 13].dropna().unique()
        return sorted(set(groups)), sorted(set(faculties))
    except Exception as e:
        print("Xatolik:", e)
        return [], []

# Tkinter oynasi
root = tk.Tk()
root.title("Nazorat Tizimi")
root.geometry("900x400")
root.state('zoomed')

# Guruh va fakultetlar
guruhlar, fakultetlar = get_groups_and_faculties_from_excel(talabalar)

# Entry va comboboxlar uchun variantlar
input_fields = {
    "Fan soati" :"",
    "Fan krediti" :"",
    "Nazorat sanasi": ""  # ← yangi qo‘shilgan qism
}


# Uqituvchi ismlarini alifbo tartibida saralash
uqituvchi_ismi = sorted(get_uqituvchi_list_from_google_sheet())

# Fanlar ro'yxatini alifbo tartibida saralash
fanlar_list = sorted(get_fanlar_from_google_sheet())

# fields dictionary-da saqlash
fields = {
    "Fakultet nomi": fakultetlar if fakultetlar else ["Fakultet 1", "Fakultet 2"],
    "Semestr": [f"{i}-semestr" for i in range(1, 9)],
    "Guruh": guruhlar if guruhlar else ["Guruh 1", "Guruh 2"],
    "Fan": fanlar_list,
    "Fan o'qituvchilari": uqituvchi_ismi,
    "Nazorat turi": ["1-ON(max-15 ball)","2-ON(max-15 ball)","Oraliq(max-30 ball)", "Yakuniy(max-70 ball)","Umumiy(max-100 ball)"],
    "Nazorat mas’uli": uqituvchi_ismi,
    "Nazorat shakli": ["1", "2", "3"],
    "Dekan nomi": uqituvchi_ismi,
    "Kafedra mudiri nomi": uqituvchi_ismi  # Shu yerda ishlatiladi
}

# print_word() funksiyasi ichida — nazorat_tur_full olinganidan keyin quyidagilarni qo‘shing:
max_ball_map = {
    "1-ON(max-15 ball)": 15,
    "2-ON(max-15 ball)": 15,
    "Oraliq(max-30 ball)": 30,
    "Yakuniy(max-70 ball)": 70,
    "Umumiy(max-100 ball)": 100
}
# max_ball = max_ball_map.get(data.get("Nazorat turi", ""), 100)  # Default 100



def calculate_baho(ball, max_ball):
    try:
        ball = float(ball)
        if ball <= 0:
            return "Kelmadi"
        if ball > max_ball:
            return "Noto‘g‘ri"
        
        baho = (ball * 5) / max_ball
        
        if baho >= 4.5:
            return "5"
        elif baho >= 3.5:
            return "4"
        elif baho >= 3.0:
            return "3"
        else:
            return "2"
    except:
        return "Noto‘g‘ri"







# Maydonlar uchun StringVar
def create_uppercase_var():
    var = tk.StringVar()
    def callback(*args):
        value = var.get()
        var.set(value.upper())
    var.trace_add('write', callback)
    return var

# Frame ichida joylashtirish
# Frame ichida joylashtirish
form_frame = tk.Frame(root)
form_frame.pack(pady=10)




comboboxes = {}
input_entries = {}
input_vars = {}




# Placeholder qo‘shuvchi funksiya Entry uchun
def set_entry_placeholder(entry, placeholder_text):
    entry.insert(0, placeholder_text)
    entry.config(fg='grey')
    def on_focus_in(event):
        if entry.get() == placeholder_text:
            entry.delete(0, tk.END)
            entry.config(fg='black')
    def on_focus_out(event):
        if not entry.get():
            entry.insert(0, placeholder_text)
            entry.config(fg='grey')
    entry.bind("<FocusIn>", on_focus_in)
    entry.bind("<FocusOut>", on_focus_out)

# Barcha maydonlarni 3 ustunli gridga joylashtirish





class SearchableCombobox(ttk.Combobox):
    def __init__(self, parent, values, **kwargs):
        super().__init__(parent, values=values, state='normal', **kwargs)
        self.values = values
        self.original_values = values
        self.bind('<KeyRelease>', self.on_keyrelease_uppercase)
        self.bind('<FocusIn>', self.on_focus_in)
        self.set("Tanlang...")

    def on_focus_in(self, event):
        if self.get() == "Tanlang...":
            self.delete(0, tk.END)

    def on_keyrelease_uppercase(self, event):
        if event.keysym in ("Up", "Down", "Left", "Right", "Return", "Tab", "Shift_L", "Shift_R"):
            return
        
        # 🔠 Kirilgan matnni katta harfga o‘zgartirish
        current = self.get().upper()
        self.delete(0, tk.END)
        self.insert(0, current)

        # 🔍 Filtrlash
        search_term = current.lower()
        current_position = self.index(tk.INSERT)

        if not search_term:
            self['values'] = self.original_values
        else:
            filtered = [item for item in self.original_values if search_term in item.lower()]
            self['values'] = filtered

        self.icursor(current_position)


import re
from datetime import datetime

def format_date_entry(event):
    widget = event.widget
    value = widget.get()

    # Ajratkichlarni aniqlash va almashtirish
    cleaned = re.sub(r"[-/\\,\s]", ".", value)

    # Turli formatdagi kirishlarni tekshirib, standart formatga o'tkazish
    for fmt in ("%d.%m.%Y", "%Y.%m.%d", "%d.%m.%y"):
        try:
            dt = datetime.strptime(cleaned, fmt)
            widget.delete(0, tk.END)
            widget.insert(0, dt.strftime("%d.%m.%Y"))
            return
        except ValueError:
            continue
    # Agar format noto‘g‘ri bo‘lsa, o‘z holida qoldiriladi




# Barcha maydonlarni 3 ustunli gridga joylashtirish
all_fields = list(fields.items()) + list(input_fields.items())
columns = 3
for index, (label, options) in enumerate(all_fields):
    row = index // columns
    col = index % columns

    tk.Label(form_frame, text=label).grid(row=row * 2, column=col, padx=10, pady=5, sticky="w")

    if isinstance(options, list):  # Searchable Combobox
        cb = SearchableCombobox(form_frame, values=options, width=50)
        cb.grid(row=row * 2 + 1, column=col, padx=10, pady=5, sticky="we")
        comboboxes[label] = cb
    else:  # Entry yoki DateEntry
        if label == "Nazorat sanasi":
            date_entry = DateEntry(form_frame, date_pattern='dd.mm.yyyy', width=47, background='darkblue', foreground='white', borderwidth=2)
            date_entry.grid(row=row * 2 + 1, column=col, padx=10, pady=5, sticky="we")
            date_entry.bind("<FocusOut>", format_date_entry)  # ← bu yerda formatlovchi funksiya ulanadi
            input_entries[label] = date_entry
        else:
            var = create_uppercase_var()
            entry = tk.Entry(form_frame, textvariable=var)
            entry.grid(row=row * 2 + 1, column=col, padx=10, pady=5, sticky="we")
            set_entry_placeholder(entry, "")
            input_vars[label] = var
            input_entries[label] = entry








# Talabalar uchun chiqadigan frame (global holatda)
# Talabalar uchun chiqadigan frame (global holatda)
scroll_canvas = tk.Canvas(root, height=400)
scrollbar = tk.Scrollbar(root, orient="vertical", command=scroll_canvas.yview)

scroll_canvas.configure(yscrollcommand=scrollbar.set)

# scrollable_frame ni scroll_canvas ichiga joylashtirish
scrollable_frame = tk.Frame(scroll_canvas)
scroll_window = scroll_canvas.create_window((0, 0), window=scrollable_frame, anchor="n")

# O‘rtaga joylashtirish uchun canvasni kengaytirishda anchor="n" o‘zgaradi
def resize_canvas(event):
    canvas_width = event.width
    scroll_canvas.itemconfig(scroll_window, width=canvas_width)

scroll_canvas.bind("<Configure>", resize_canvas)

# scrollable_frame ni o'zgarishlarini aniqlash
scrollable_frame.bind(
    "<Configure>",
    lambda e: scroll_canvas.configure(scrollregion=scroll_canvas.bbox("all"))
)

scroll_canvas.pack(fill='both', expand=True, padx=20, pady=10)
scrollbar.pack(side="right", fill="y")
scroll_canvas.pack_forget()  # Boshlanishda yashirin




def _on_mousewheel(event):
    if event.delta:
        scroll_canvas.yview_scroll(-1 * int(event.delta / 120), "units")
    else:
        scroll_canvas.yview_scroll(1 if event.num == 5 else -1, "units")

# Barcha joyda ishlashi uchun — har qanday joyda scroll bo‘lishi uchun
root.bind_all("<MouseWheel>", _on_mousewheel)        # Windows va macOS
root.bind_all("<Button-4>", _on_mousewheel)          # Linux: scroll up
root.bind_all("<Button-5>", _on_mousewheel)          # Linux: scroll down



def saqlash():
    xatolik_label = tk.Label(root, text="", fg="red", font=("Arial", 12))
    xatolik_label.pack()
    xatolik_label.pack_forget()  # Boshlanishda yashirin bo‘ladi

    global ball_entries
    ball_entries = {}  # Har saqlashda tozalanadi

    for widget in scrollable_frame.winfo_children():
        widget.destroy()

    data = {label: entry.get() for label, entry in input_entries.items()}
    data.update({label: cb.get() for label, cb in comboboxes.items()})

    tanlangan_guruh = data.get("Guruh")
    if not tanlangan_guruh:
        warning_frame = tk.Frame(scrollable_frame)
        warning_frame.pack(pady=10, fill='x')
        tk.Label(warning_frame, text="❌ Iltimos, guruhni tanlang.", fg="red", font=("Arial", 12)).pack(anchor="center")
        scroll_canvas.pack(fill='both', expand=True, padx=20, pady=10)
        return

    try:
        df = pd.read_excel(talabalar)
        mos_talabalar = df[df['Guruh'] == tanlangan_guruh]
        sorted_talabalar = mos_talabalar[['Talaba ID', 'To‘liq ismi']].dropna().sort_values(by='To‘liq ismi')

        if sorted_talabalar.empty:
            empty_frame = tk.Frame(scrollable_frame)
            empty_frame.pack(pady=10, fill='x')
            tk.Label(empty_frame, text=f"❌ '{tanlangan_guruh}' guruhiga mos talaba topilmadi.",
                     fg="red", font=("Arial", 12)).pack(anchor="center")
        else:
            tk.Label(scrollable_frame, text=f"✅ '{tanlangan_guruh}' guruhidagi talabalar soni: {len(sorted_talabalar)} nafar", font=("Arial", 11, "bold")).pack(pady=10, anchor="center")

            header_frame = tk.Frame(scrollable_frame)
            header_frame.pack(pady=(5, 2))
            tk.Label(header_frame, text="T/R", width=5, font=("Arial", 10, "bold")).pack(side="left")
            tk.Label(header_frame, text="F.I.O", width=40, font=("Arial", 10, "bold")).pack(side="left")
            tk.Label(header_frame, text="Hemis ID", width=20, font=("Arial", 10, "bold")).pack(side="left")
            tk.Label(header_frame, text="Ball", width=10, font=("Arial", 10, "bold")).pack(side="left")

            for idx, (_, row) in enumerate(sorted_talabalar.iterrows(), start=1):
                row_frame = tk.Frame(scrollable_frame)
                row_frame.pack(pady=2)

                row_frame_inner = tk.Frame(row_frame)
                row_frame_inner.pack(anchor="center")

                tk.Label(row_frame_inner, text=f"{idx}", width=5, anchor="w").pack(side="left")
                tk.Label(row_frame_inner, text=f"{row['To‘liq ismi']}", width=40, anchor="w", wraplength=500).pack(side="left")
                tk.Label(row_frame_inner, text=f"{row['Talaba ID']}", width=20, anchor="w").pack(side="left")

                entry = tk.Entry(row_frame_inner, width=10)
                entry.pack(side="left", padx=10)
                ball_entries[row['Talaba ID']] = entry

    except Exception as e:
        tk.Label(scrollable_frame, text="❌ Faylni o‘qishda xatolik: " + str(e), fg="red").pack()

    scroll_canvas.pack(fill='both', expand=True, padx=20, pady=10)







from docx import Document
from docx.shared import Pt

def replace_text_in_doc(doc, replace_map):
    for p in doc.paragraphs:
        full_text = ''.join(run.text for run in p.runs)
        for key, value in replace_map.items():
            if key in full_text:
                full_text = full_text.replace(key, value)
        for i, run in enumerate(p.runs):
            run.text = ''
        if p.runs:
            p.runs[0].text = full_text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    full_text = ''.join(run.text for run in p.runs)
                    for key, value in replace_map.items():
                        if key in full_text:
                            full_text = full_text.replace(key, value)
                    for i, run in enumerate(p.runs):
                        run.text = ''
                    if p.runs:
                        p.runs[0].text = full_text
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # ⬅️ Import qilish kerak

def print_word():
    data = {label: entry.get() for label, entry in input_entries.items()}
    data.update({label: cb.get() for label, cb in comboboxes.items()})

    alo_5 = 0
    yaxshi_4 = 0
    qoniqarli_3 = 0
    qoniqarsiz_2 = 0
    kelmadi = 0

    tanlangan_guruh = data.get("Guruh")
    if not tanlangan_guruh:
        print("❌ Guruh tanlanmagan.")
        return

    try:
        df = pd.read_excel(talabalar)
        mos_talabalar = df[df['Guruh'] == tanlangan_guruh]
        sorted_talabalar = mos_talabalar[['Talaba ID', 'To‘liq ismi']].dropna(subset=['Talaba ID', 'To‘liq ismi'])
    except Exception as e:
        print("❌ Talabalarni o‘qishda xatolik:", e)
        return

    try:
        doc_name = create_word()
        doc = Document("docx_temp.docx")
    except Exception as e:
        print("❌ docx_temp.docx ochishda xatolik:", e)
        return

    import re
    nazorat_tur_full = data.get("Nazorat turi", "")
    nazorat_tur = re.sub(r"\(.*?\)", "", nazorat_tur_full).strip()

    max_ball_map = {
        "1-ON(max-15 ball)": 15,
        "2-ON(max-15 ball)": 15,
        "Oraliq(max-30 ball)": 30,
        "Yakuniy(max-70 ball)": 70,
        "Umumiy(max-100 ball)": 100
    }
    max_ball = max_ball_map.get(nazorat_tur_full, 100)

    replace_map = {
        "{fakultet}": data.get("Fakultet nomi", ""),
        "{semester}": data.get("Semestr", ""),
        "{guruh}": data.get("Guruh", ""),
        "{fan}": data.get("Fan", ""),
        "{fan_uqituvchi}": data.get("Fan o'qituvchilari", ""),
        "{nazorat_turi}": nazorat_tur,
        "{nazorat_tur}dan to‘plagan ballar": f"{nazorat_tur}dan to‘plagan ballar",
        "{nazorat_masuli}": data.get("Nazorat mas’uli", ""),
        "{soat}": data.get("Fan soati", ""),
        "{kredit}": data.get("Fan krediti", ""),
        "{nazorat_sanasi}": data.get("Nazorat sanasi", ""),
        "{tur}": data.get("Nazorat shakli", ""),
        "{dekan}": data.get("Dekan nomi", ""),
        "{mudir}": data.get("Kafedra mudiri nomi", "")
    }

    for table in doc.tables:
        if any("{tr}" in cell.text for row in table.rows for cell in row.cells):
            for i, row in enumerate(table.rows):
                if "{tr}" in row.cells[0].text:
                    table._tbl.remove(row._tr)
                    break

        for idx, (talaba_id, entry_widget) in enumerate(ball_entries.items(), start=1):
            parent_widgets = entry_widget.master.winfo_children()
            talaba_ismi = parent_widgets[1].cget("text") if len(parent_widgets) > 1 else ""
            ball = entry_widget.get().strip()

            ball_text = "Noto‘g‘ri"
            baho_text = "Noto‘g‘ri"

            if not ball or ball == "0":
                ball_text = "Kelmadi"
                baho_text = "Kelmadi"
                kelmadi += 1
            else:
                try:
                    ball_val = float(ball)
                    if ball_val > max_ball:
                        messagebox.showerror("Xatolik", f"❌ {talaba_ismi} uchun kiritilgan ball ({ball_val}) maksimal balldan ({max_ball}) katta.")
                        return
                    if ball_val < 0:
                        messagebox.showerror("Xatolik", f"❌ {talaba_ismi} uchun kiritilgan ball ({ball_val}) 0 dan kichik.")
                        return
                    ball_text = str(int(ball_val)) if ball_val.is_integer() else str(ball_val)
                    baho_text = calculate_baho(ball_val, max_ball)

                    if baho_text == "5":
                        alo_5 += 1
                    elif baho_text == "4":
                        yaxshi_4 += 1
                    elif baho_text == "3":
                        qoniqarli_3 += 1
                    elif baho_text == "2":
                        qoniqarsiz_2 += 1
                    elif baho_text.lower() == "kelmadi":
                        kelmadi += 1
                except ValueError:
                    ball_text = "Noto‘g‘ri"
                    baho_text = "Noto‘g‘ri"


                except ValueError:
                    messagebox.showerror("Xatolik", f"❌ {talaba_ismi} uchun noto‘g‘ri ball qiymati: {ball}")
                    return

            new_row = table.add_row()
            if len(new_row.cells) < 6:
                print(f"❌ Jadval ustunlari yetarli emas. Topilgan: {len(new_row.cells)}")
                continue

            new_row.cells[0].text = str(idx)
            new_row.cells[1].text = talaba_ismi
            new_row.cells[2].text = str(talaba_id)
            new_row.cells[3].text = ball_text
            new_row.cells[4].text = baho_text
            new_row.cells[5].text = ""

            for cell_index, cell in enumerate(new_row.cells):
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if cell_index in [3, 4] else WD_PARAGRAPH_ALIGNMENT.LEFT
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

    replace_map.update({
        "{student_soni}": str(len(ball_entries)),
        "{alo_5}": str(int(alo_5/2)),
        "{yaxshi_4}": str(int(yaxshi_4/2)),
        "{qoniqarli_3}": str(int(qoniqarli_3/2)),
        "{qoniqarsiz_2}": str(int(qoniqarsiz_2/2)),
        "{kelmadi}": str(int(kelmadi/2))
    })


    replace_text_in_doc(doc, replace_map)

    fan_nomi = data.get("Fan", "").replace(" ", "_")
    semester = data.get("Semestr", "").replace(" ", "_")
    nazorat_turi = data.get("Nazorat turi", "").replace(" ", "_")
    nazorat_shakli = data.get("Nazorat shakli", "").replace(" ", "_")

    filename = f"{tanlangan_guruh}_{semester}_{fan_nomi}_{nazorat_turi}_{nazorat_shakli}_tur.docx"
    output_path = os.path.join("data", filename)
    doc.save(output_path)

    delete_file("docx_temp.docx")

    print(f"✅ Word hujjat yaratildi: {output_path}")
    send_file_to_telegram_group(output_path)

import requests

def send_file_to_telegram_group(file_path):
    bot_token = "7988992200:AAHguvG_iPE6ZG6gdsCjWKJ17fv6vpPzJsQ"
    chat_id = "-1002635664827"  # Guruh ID
    url = f"https://api.telegram.org/bot{bot_token}/sendDocument"

    with open(file_path, "rb") as f:
        files = {"document": f}
        data = {"chat_id": chat_id}
        response = requests.post(url, data=data, files=files)

    if response.status_code == 200:
        print("✅ Hujjat data papkaga yuborildi!")
    else:
        print("❌ Hujjat Hujjat data papkaga yuborishda xatolik:", response.text)

# Word hujjat yaratib bo‘lgandan so‘ng chaqirasiz:












# Tugmalar uchun frame
button_frame = tk.Frame(root)
button_frame.pack(pady=10)

# Saqlash tugmasi
tk.Button(button_frame, text="✅ Saqlash", command=saqlash, bg="#4CAF50", fg="white", font=("Arial", 11, "bold")).pack(side="left", padx=10)

# Chop etish tugmasi (hozircha shu ham saqlash funksiyasini chaqiradi)
tk.Button(button_frame, text="✅ Chop etish", command=print_word, bg="#4CAF50", fg="white", font=("Arial", 11, "bold")).pack(side="left", padx=10)

root.mainloop()
