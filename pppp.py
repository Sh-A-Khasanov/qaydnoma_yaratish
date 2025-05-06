import tkinter as tk
from tkinter import ttk, messagebox
from tkcalendar import DateEntry
import pandas as pd
import glob
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import requests
from io import StringIO
from datetime import datetime
import re
from word_yasash import create_word, delete_file

# Colors
TEXT_COLOR = "#FFFFFF"
BORDER_COLOR = "#012C6E"
BUTTON_COLOR = "#04AA6D"
BG_COLOR = "#F5F6F5"  # Light gray background
SECONDARY_BG = "#E8ECEF"  # Secondary background

# Font settings
FONT_FAMILY = "Arial"
LABEL_FONT = (FONT_FAMILY, 12, "bold")
ENTRY_FONT = (FONT_FAMILY, 11)
BUTTON_FONT = (FONT_FAMILY, 11, "bold")

# Translation dictionaries
translations = {
    "uz": {
        "fields": {
            "Fakultet nomi": "Fakultet nomi",
            "Semestr": "Semestr",
            "Guruh": "Guruh",
            "Fan": "Fan",
            "Fan o'qituvchilari": "Fan o'qituvchilari",
            "Nazorat shakli": "Nazorat shakli",
            "Nazorat mas’uli": "Nazorat mas’uli",
            "Nazorat turi": "Nazorat turi",
            "Dekan nomi": "Dekan nomi",
            "Kafedra mudiri nomi": "Kafedra mudiri nomi"
        },
        "input_fields": {
            "Fan soati": "Fan soati",
            "Fan krediti": "Fan krediti",
            "Nazorat sanasi": "Nazorat sanasi"
        },
        "buttons": {
            "Saqlash": "✅ Saqlash",
            "Chop etish": "🖨 Chop etish"
        },
        "placeholder": "Tanlang...",
        "messages": {
            "student_count": "✅ '{group}' guruhidagi talabalar soni: {count} nafar"
        }
    },
    "ru": {
        "fields": {
            "Fakultet nomi": "Название факультета",
            "Semestr": "Семестр",
            "Guruh": "Группа",
            "Fan": "Предмет",
            "Fan o'qituvchilari": "Преподаватели предмета",
            "Nazorat shakli": "Форма контроля",
            "Nazorat mas’uli": "Ответственный за контроль",
            "Nazorat turi": "Тип контроля",
            "Dekan nomi": "Имя декана",
            "Kafedra mudiri nomi": "Имя заведующего кафедрой"
        },
        "input_fields": {
            "Fan soati": "Часы предмета",
            "Fan krediti": "Кредиты предмета",
            "Nazorat sanasi": "Дата контроля"
        },
        "buttons": {
            "Saqlash": "✅ Сохранить",
            "Chop etish": "🖨 Печать"
        },
        "placeholder": "Выберите...",
        "messages": {
            "student_count": "✅ Количество студентов в группе '{group}': {count} чел."
        }
    },
    "en": {
        "fields": {
            "Fakultet nomi": "Faculty Name",
            "Semestr": "Semester",
            "Guruh": "Group",
            "Fan": "Subject",
            "Fan o'qituvchilari": "Subject Teachers",
            "Nazorat shakli": "Control Form",
            "Nazorat mas’uli": "Control Responsible",
            "Nazorat turi": "Control Type",
            "Dekan nomi": "Dean’s Name",
            "Kafedra mudiri nomi": "Head of Department Name"
        },
        "input_fields": {
            "Fan soati": "Subject Hours",
            "Fan krediti": "Subject Credits",
            "Nazorat sanasi": "Control Date"
        },
        "buttons": {
            "Saqlash": "✅ Save",
            "Chop etish": "🖨 Print"
        },
        "placeholder": "Select...",
        "messages": {
            "student_count": "✅ Number of students in group '{group}': {count}"
        }
    }
}

# Control form options translation
control_form_translations = {
    "uz": [
        "1-ON(max-15 ball)",
        "2-ON(max-15 ball)",
        "Oraliq(max-30 ball)",
        "Yakuniy(max-70 ball)",
        "Umumiy(max-100 ball)"
    ],
    "ru": [
        "1-ОН(макс-15 баллов)",
        "2-ОН(макс-15 баллов)",
        "Oraliq(макс-30 баллов)",
        "Yakuniy(макс-70 баллов)",
        "Umumiy(макс-100 баллов)"
    ],
    "en": [
        "1-ON(max-15 points)",
        "2-ON(max-15 points)",
        "Oraliq(max-30 points)",
        "Yakuniy(max-70 points)",
        "Umumiy(max-100 points)"
    ]
}

# Global variable for selected language
selected_language = "uz"  # Default to Uzbek

# Data fetching functions
def get_code_from_google_sheet(entered_code):
    sheet_id = "1eJ6LDB61vZ8ZW2IAyseKOnLigUHEbas6F0bwTquqIeU"
    sheet_name = "KOD"
    url = f"https://docs.google.com/spreadsheets/d/{sheet_id}/gviz/tq?tqx=out:csv&sheet={sheet_name}"
    try:
        response = requests.get(url)
        response.raise_for_status()
        df = pd.read_csv(StringIO(response.text), header=None)  # Ustun nomlarini o'qimaslik uchun header=None
        if not df.empty and df.shape[1] >= 2:  # Kamida 2 ustun borligini tekshirish
            # A ustuni (0-indeks) bo'yicha kiritilgan kodni qidirish
            for index, value in df[0].items():
                if str(value).strip().lower() == str(entered_code).strip().lower():
                    # Agar kod topilsa, B ustunidagi (1-indeks) qiymatni chop etish
                    corresponding_value = df.iloc[index, 1]
                    if pd.notna(corresponding_value):  # Qiymat NaN emasligini tekshirish
                        print(f"✅ Topilgan kod: {entered_code}, B ustunidagi qiymat: {corresponding_value}")
                        return str(corresponding_value).strip()
                    else:
                        print(f"❌ Kod topildi ({entered_code}), lekin B ustunidagi qiymat bo'sh.")
                        return None
            print(f"❌ Kiritilgan kod ({entered_code}) mos kelmadi")
            return None
        else:
            print("❌ 'KOD' sahifasida ma'lumotlar yetarli emas yoki topilmadi.")
            return None
    except Exception as e:
        print(f"❌ Google Sheet'dan kod olishda xatolik: {e}")
        return None



def get_talabalar_from_google_sheet(sheet_name):
    if not sheet_name:
        print("❌ Sahifa nomi kiritilmadi!")
        return None
    sheet_id = "1eJ6LDB61vZ8ZW2IAyseKOnLigUHEbas6F0bwTquqIeU"
    url = f"https://docs.google.com/spreadsheets/d/{sheet_id}/gviz/tq?tqx=out:csv&sheet={sheet_name}"
    try:
        response = requests.get(url)
        response.raise_for_status()
        df = pd.read_csv(StringIO(response.text))
        return df
    except Exception as e:
        print(f"❌ Google Sheet'dan talabalar ro'yxatini olishda xatolik (sahifa: {sheet_name}): {e}")
        return None
    



def get_uqituvchi_list_from_google_sheet():
    sheet_id = "1eJ6LDB61vZ8ZW2IAyseKOnLigUHEbas6F0bwTquqIeU"
    sheet_name = "Xodimlar"
    url = f"https://docs.google.com/spreadsheets/d/{sheet_id}/gviz/tq?tqx=out:csv&sheet={sheet_name}"
    try:
        response = requests.get(url)
        response.raise_for_status()
        df = pd.read_csv(StringIO(response.text))
        if {'Ismi', 'Familiya', "Otasining ismi"}.issubset(df.columns):
            full_names = df[['Familiya', 'Ismi', 'Otasining ismi']].fillna('').astype(str).apply(
                lambda x: f"{x['Familiya']} {x['Ismi']} {x['Otasining ismi']}".strip(), axis=1)
            return list(set(full_names))
        else:
            print("❌ Kutilgan ustun nomlari topilmadi.")
            return []
    except Exception as e:
        print("❌ Google Sheet dan o‘qishda xatolik:", e)
        return []

def get_fanlar_from_google_sheet():
    sheet_id = "1eJ6LDB61vZ8ZW2IAyseKOnLigUHEbas6F0bwTquqIeU"
    sheet_name = "Fanlar ro'yxati"
    url = f"https://docs.google.com/spreadsheets/d/{sheet_id}/gviz/tq?tqx=out:csv&sheet={sheet_name}"
    try:
        response = requests.get(url)
        response.raise_for_status()
        df = pd.read_csv(StringIO(response.text))
        if 'Nomi' in df.columns:
            fanlar = df['Nomi'].dropna().astype(str).tolist()
            return fanlar
        else:
            print("❌ 'Nomi' ustuni topilmadi.")
            return []
    except Exception as e:
        print("❌ Fanlar listidan o‘qishda xatolik:", e)
        return []

def get_groups_and_faculties_from_google_sheet(sheet_name):
    df = get_talabalar_from_google_sheet(sheet_name)
    if df is not None:
        try:
            groups = df['Guruh'].dropna().unique()
            faculties = df['Fakultet'].dropna().unique()
            return sorted(set(groups)), sorted(set(faculties))
        except Exception as e:
            print(f"❌ Guruhlar va fakultetlarni olishda xatolik: {e}")
            return [], []
    return [], []




# Tkinter window
root = tk.Tk()
root.title("Qaydnoma shakillantirish")
root.geometry("1200x600")
root.state('zoomed')
root.configure(bg=BG_COLOR)

# Data initialization
# Data initialization
guruhlar = []
fakultetlar = []
uqituvchi_ismi = sorted(get_uqituvchi_list_from_google_sheet())
fanlar_list = sorted(get_fanlar_from_google_sheet())

fields = {
    "Fakultet nomi": fakultetlar if fakultetlar else ["Fakultet 1", "Fakultet 2"],
    "Semestr": [f"{i}-semestr" for i in range(1, 13)],
    "Guruh": guruhlar if guruhlar else ["Guruh 1", "Guruh 2"],
    "Fan": fanlar_list,
    "Fan o'qituvchilari": uqituvchi_ismi,
    "Nazorat shakli": ["1-ON(max-15 ball)", "2-ON(max-15 ball)", "Oraliq(max-30 ball)", "Yakuniy(max-70 ball)", "Umumiy(max-100 ball)"],
    "Nazorat mas’uli": uqituvchi_ismi,
    "Nazorat turi": ["1", "2", "3"],
    "Dekan nomi": uqituvchi_ismi,
    "Kafedra mudiri nomi": uqituvchi_ismi
}

input_fields = {
    "Fan soati": "",
    "Fan krediti": "",
    "Nazorat sanasi": ""
}

# Control form mappings
nazorat_shakli_map = {
    "1": "1",
    "2": "1a",
    "3": "1b"
}
max_ball_map = {
    "1-ON(max-15 ball)": 15,
    "2-ON(max-15 ball)": 15,
    "Oraliq(max-30 ball)": 30,
    "Yakuniy(max-70 ball)": 70,
    "Umumiy(max-100 ball)": 100,
    "1-ОН(макс-15 баллов)": 15,
    "2-ОН(макс-15 баллов)": 15,
    "Промежуточный(макс-30 баллов)": 30,
    "Итоговый(макс-70 баллов)": 70,
    "Общий(макс-100 баллов)": 100,
    "1-ON(max-15 points)": 15,
    "2-ON(max-15 points)": 15,
    "Midterm(max-30 points)": 30,
    "Final(max-70 points)": 70,
    "Total(max-100 points)": 100
}

def calculate_baho(ball, max_ball):
    try:
        ball = float(ball)
        if ball <= 0:
            return "Kelmadi"
        if ball > max_ball:
            return "Noto‘g‘ri"
        baho = (ball * 5) / max_ball
        if baho >= 4.5:
            return "5"
        elif baho >= 3.5:
            return "4"
        elif baho >= 3.0:
            return "3"
        else:
            return "O'tmadi"
    except:
        return "Noto‘g‘ri"

def create_uppercase_var():
    var = tk.StringVar()
    def callback(*args):
        value = var.get()
        var.set(value.upper())
    var.trace_add('write', callback)
    return var

# Form frame
form_frame = tk.Frame(root, bg=BG_COLOR, bd=2, relief="groove")
form_frame.pack(pady=20, padx=20, fill="x")

# Button frame
button_frame = tk.Frame(root, bg=TEXT_COLOR)
button_frame.pack(pady=20)

comboboxes = {}
input_entries = {}
input_vars = {}

# Language selection
def select_language():
    global selected_language
    lang_window = tk.Toplevel(root)
    lang_window.title("Avtorizatsiya")
    lang_window.geometry("400x300")
    lang_window.configure(bg=BG_COLOR)
    lang_window.grab_set()

    # Kod kiritish uchun yorliq va maydon
    tk.Label(lang_window, text="Kodni kiriting / Введите код / Enter the code", bg=BG_COLOR, font=LABEL_FONT).pack(pady=10)
    code_entry = tk.Entry(lang_window, show="*", width=20, font=ENTRY_FONT, bg=SECONDARY_BG, fg=BORDER_COLOR, bd=2, relief="groove")
    code_entry.pack(pady=5)

    # Xato xabari uchun yorliq
    error_label = tk.Label(lang_window, text="", fg="red", bg=BG_COLOR, font=("Arial", 10))
    error_label.pack(pady=5)

    # Til tanlash uchun yorliq
    tk.Label(lang_window, text="Tilni tanlang / Выберите язык / Select Language", bg=BG_COLOR, font=LABEL_FONT).pack(pady=10)

    def check_code_and_set_language(lang):
        entered_code = code_entry.get().strip()
        global corresponding_value
        corresponding_value = get_code_from_google_sheet(entered_code)
        if corresponding_value:
            global selected_language, guruhlar, fakultetlar
            selected_language = lang
            # corresponding_value ni sheet_name sifatida ishlatish
            guruhlar, fakultetlar = get_groups_and_faculties_from_google_sheet(sheet_name=corresponding_value)
            if guruhlar or fakultetlar:
                print(f"✅ {corresponding_value} sahifasidan guruhlar va fakultetlar muvaffaqiyatli olindi.")
                lang_window.destroy()
                initialize_ui()
            else:
                error_label.config(text=f"❌ {corresponding_value} sahifasidan guruhlar va fakultetlar olishda xato!")
        else:
            error_label.config(text=f"❌ Kiritilgan kod ({entered_code}) xato.")



    # Til tanlash tugmalari
    for lang, label in [("uz", "O'zbek"), ("ru", "Русский"), ("en", "English")]:
        tk.Button(lang_window, text=label, bg=BUTTON_COLOR, fg=TEXT_COLOR, font=BUTTON_FONT,
                  command=lambda l=lang: check_code_and_set_language(l)).pack(pady=5, fill="x", padx=20)



def set_language(lang, window):
    global selected_language
    selected_language = lang
    window.destroy()
    initialize_ui()

def initialize_ui():
    # Clear existing widgets
    for widget in form_frame.winfo_children():
        widget.destroy()
    for widget in button_frame.winfo_children():
        widget.destroy()

    # Fields ni yangilash
    fields = {
        translations[selected_language]["fields"]["Fakultet nomi"]: fakultetlar if fakultetlar else ["Fakultet 1", "Fakultet 2"],
        translations[selected_language]["fields"]["Semestr"]: [f"{i}-semestr" for i in range(1, 13)],
        translations[selected_language]["fields"]["Guruh"]: guruhlar if guruhlar else ["Guruh 1", "Guruh 2"],
        translations[selected_language]["fields"]["Fan"]: fanlar_list,
        translations[selected_language]["fields"]["Fan o'qituvchilari"]: uqituvchi_ismi,
        translations[selected_language]["fields"]["Nazorat shakli"]: control_form_translations[selected_language],
        translations[selected_language]["fields"]["Nazorat mas’uli"]: uqituvchi_ismi,
        translations[selected_language]["fields"]["Nazorat turi"]: ["1", "2", "3"],
        translations[selected_language]["fields"]["Dekan nomi"]: uqituvchi_ismi,
        translations[selected_language]["fields"]["Kafedra mudiri nomi"]: uqituvchi_ismi
    }

    # Translate fields
    translated_fields = {
        key: value for key, value in fields.items()
    }

    translated_input_fields = {
        translations[selected_language]["input_fields"][key]: value
        for key, value in input_fields.items()
    }

    # Place fields
    all_fields = list(translated_fields.items()) + list(translated_input_fields.items())
    columns = 3
    for index, (label, options) in enumerate(all_fields):
        row = index // columns
        col = index % columns

        tk.Label(form_frame, text=label, font=LABEL_FONT, bg=BG_COLOR, fg=BORDER_COLOR).grid(row=row * 2, column=col, padx=10, pady=5, sticky="w")

        if isinstance(options, list):
            cb = SearchableCombobox(form_frame, values=options, width=40)
            cb.grid(row=row * 2 + 1, column=col, padx=10, pady=5, sticky="we")
            comboboxes[label] = cb
        else:
            if label == translations[selected_language]["input_fields"]["Nazorat sanasi"]:
                date_entry = DateEntry(form_frame, date_pattern='dd.mm.yyyy', width=47, background=BORDER_COLOR, foreground=TEXT_COLOR, borderwidth=2, font=ENTRY_FONT)
                date_entry.grid(row=row * 2 + 1, column=col, padx=10, pady=5, sticky="we")
                date_entry.bind("<FocusOut>", format_date_entry)
                input_entries[label] = date_entry
            else:
                var = create_uppercase_var()
                entry = tk.Entry(form_frame, textvariable=var, bg=SECONDARY_BG, fg=BORDER_COLOR, font=ENTRY_FONT, bd=2, relief="groove")
                entry.grid(row=row * 2 + 1, column=col, padx=10, pady=5, sticky="we")
                set_entry_placeholder(entry, "")
                input_vars[label] = var
                input_entries[label] = entry

    # Create buttons
    create_button(button_frame, translations[selected_language]["buttons"]["Saqlash"], saqlash)
    create_button(button_frame, translations[selected_language]["buttons"]["Chop etish"], print_word)


def set_entry_placeholder(entry, placeholder_text):
    entry.insert(0, placeholder_text)
    entry.config(fg='grey', font=ENTRY_FONT)
    def on_focus_in(event):
        if entry.get() == placeholder_text:
            entry.delete(0, tk.END)
            entry.config(fg=BORDER_COLOR)
    def on_focus_out(event):
        if not entry.get():
            entry.insert(0, placeholder_text)
            entry.config(fg='grey')
    entry.bind("<FocusIn>", on_focus_in)
    entry.bind("<FocusOut>", on_focus_out)

class SearchableCombobox(ttk.Combobox):
    def __init__(self, parent, values, **kwargs):
        super().__init__(parent, values=values, state='normal', font=ENTRY_FONT, **kwargs)
        self.values = values
        self.original_values = values
        self.bind('<KeyRelease>', self.on_keyrelease_uppercase)
        self.bind('<FocusIn>', self.on_focus_in)
        self.set(translations[selected_language]["placeholder"])
        self.configure(style="Custom.TCombobox")

    def on_focus_in(self, event):
        if self.get() == translations[selected_language]["placeholder"]:
            self.delete(0, tk.END)

    def on_keyrelease_uppercase(self, event):
        if event.keysym in ("Up", "Down", "Left", "Right", "Return", "Tab", "Shift_L", "Shift_R"):
            return
        current = self.get().upper()
        self.delete(0, tk.END)
        self.insert(0, current)
        search_term = current.lower()
        current_position = self.index(tk.INSERT)
        if not search_term:
            self['values'] = self.original_values
        else:
            filtered = [item for item in self.original_values if search_term in item.lower()]
            self['values'] = filtered
        self.icursor(current_position)

def format_date_entry(event):
    widget = event.widget
    value = widget.get()
    cleaned = re.sub(r"[-/\\,\s]", ".", value)
    for fmt in ("%d.%m.%Y", "%Y.%m.%d", "%d.%m.%y"):
        try:
            dt = datetime.strptime(cleaned, fmt)
            widget.delete(0, tk.END)
            widget.insert(0, dt.strftime("%d.%m.%Y"))
            return
        except ValueError:
            continue

# Combobox style
style = ttk.Style()
style.theme_use('clam')
style.configure("Custom.TCombobox",
                fieldbackground=SECONDARY_BG,
                background=BORDER_COLOR,
                foreground=BORDER_COLOR,
                bordercolor=BORDER_COLOR,
                arrowcolor=BORDER_COLOR)
style.map("Custom.TCombobox",
          fieldbackground=[('readonly', SECONDARY_BG)],
          selectbackground=[('readonly', SECONDARY_BG)],
          selectforeground=[('readonly', BORDER_COLOR)])

# Scrollable frame
scroll_canvas = tk.Canvas(root, height=400, bg=BG_COLOR, highlightthickness=0)
scrollbar = tk.Scrollbar(root, orient="vertical", command=scroll_canvas.yview, bg=BORDER_COLOR, troughcolor=SECONDARY_BG)
scroll_canvas.configure(yscrollcommand=scrollbar.set)

scrollable_frame = tk.Frame(scroll_canvas, bg=BG_COLOR)
scroll_window = scroll_canvas.create_window((0, 0), window=scrollable_frame, anchor="n")

def resize_canvas(event):
    canvas_width = event.width
    scroll_canvas.itemconfig(scroll_window, width=canvas_width)

scroll_canvas.bind("<Configure>", resize_canvas)
scrollable_frame.bind("<Configure>", lambda e: scroll_canvas.configure(scrollregion=scroll_canvas.bbox("all")))
scroll_canvas.pack(fill='both', expand=True, padx=20, pady=10)
scrollbar.pack(side="right", fill="y")
scroll_canvas.pack_forget()

def _on_mousewheel(event):
    if event.delta:
        scroll_canvas.yview_scroll(-1 * int(event.delta / 120), "units")
    else:
        scroll_canvas.yview_scroll(1 if event.num == 5 else -1, "units")

root.bind_all("<MouseWheel>", _on_mousewheel)
root.bind_all("<Button-4>", _on_mousewheel)
root.bind_all("<Button-5>", _on_mousewheel)



def saqlash():
    xatolik_label = tk.Label(root, text="", fg="red", font=("Arial", 12))
    xatolik_label.pack()
    xatolik_label.pack_forget()

    global ball_entries
    ball_entries = {}

    for widget in scrollable_frame.winfo_children():
        widget.destroy()

    data = {label: entry.get() for label, entry in input_entries.items()}
    data.update({label: cb.get() for label, cb in comboboxes.items()})

    tanlangan_guruh = data.get(translations[selected_language]["fields"]["Guruh"])
    if not tanlangan_guruh or tanlangan_guruh == translations[selected_language]["placeholder"]:
        warning_frame = tk.Frame(scrollable_frame, bg=BG_COLOR)
        warning_frame.pack(pady=10, fill='x')
        tk.Label(warning_frame, text="❌ Iltimos, guruhni tanlang.", fg="red", font=("Arial", 12), bg=BG_COLOR).pack(anchor="center")
        scroll_canvas.pack(fill='both', expand=True, padx=20, pady=10)
        return

    try:
        df = get_talabalar_from_google_sheet(corresponding_value)
        if df is None:
            tk.Label(scrollable_frame, text="❌ Talabalar ro'yxatini olishda xatolik yuz berdi.", fg="red", bg=BG_COLOR).pack()
            return
        mos_talabalar = df[df['Guruh'] == tanlangan_guruh]
        sorted_talabalar = mos_talabalar[['Talaba ID', 'To‘liq ismi']].dropna().sort_values(by='To‘liq ismi')

        if sorted_talabalar.empty:
            empty_frame = tk.Frame(scrollable_frame, bg=BG_COLOR)
            empty_frame.pack(pady=10, fill='x')
            tk.Label(empty_frame, text=f"❌ '{tanlangan_guruh}' guruhiga mos talaba topilmadi.",
                    fg="red", font=("Arial", 12), bg=BG_COLOR).pack(anchor="center")
        else:
            message = translations[selected_language]["messages"]["student_count"].format(group=tanlangan_guruh, count=len(sorted_talabalar))
            tk.Label(scrollable_frame, text=message,
                    font=("Arial", 11, "bold"), bg=BG_COLOR, fg=BORDER_COLOR).pack(pady=10, anchor="center")

            header_frame = tk.Frame(scrollable_frame, bg=SECONDARY_BG)
            header_frame.pack(pady=(5, 2))

            labels = [
                ("T/R                            F.I.O               ", 30),
                ("               Hemis ID", 20),
                ("              Ball", 20)
            ]

            for idx, (text, width) in enumerate(labels):
                tk.Label(header_frame, text=text, width=width, font=("Arial", 10, "bold"), bg=SECONDARY_BG, fg=BORDER_COLOR).grid(row=0, column=idx, padx=5)

            for idx, (_, row) in enumerate(sorted_talabalar.iterrows(), start=1):
                row_frame = tk.Frame(scrollable_frame, bg=BG_COLOR)
                row_frame.pack(pady=2)

                row_frame_inner = tk.Frame(row_frame, bg=BG_COLOR)
                row_frame_inner.pack(anchor="center")

                tk.Label(row_frame_inner, text=f"{idx}", width=5, anchor="w", bg=BG_COLOR, fg=BORDER_COLOR).pack(side="left")
                tk.Label(row_frame_inner, text=f"{row['To‘liq ismi']}", width=40, anchor="w", wraplength=500, bg=BG_COLOR, fg=BORDER_COLOR).pack(side="left")
                tk.Label(row_frame_inner, text=f"{row['Talaba ID']}", width=20, anchor="w", bg=BG_COLOR, fg=BORDER_COLOR).pack(side="left")

                entry = tk.Entry(row_frame_inner, width=10, bg=SECONDARY_BG, fg=BORDER_COLOR, font=ENTRY_FONT, bd=2, relief="groove")
                entry.pack(side="left", padx=10)
                ball_entries[row['Talaba ID']] = entry

    except Exception as e:
        tk.Label(scrollable_frame, text="❌ Ma'lumotlarni o‘qishda xatolik: " + str(e), fg="red", bg=BG_COLOR).pack()

    scroll_canvas.pack(fill='both', expand=True, padx=20, pady=10)


def replace_text_in_doc(doc, replace_map):
    for p in doc.paragraphs:
        full_text = ''.join(run.text for run in p.runs)
        for key, value in replace_map.items():
            if key in full_text:
                full_text = full_text.replace(key, value)
        for i, run in enumerate(p.runs):
            run.text = ''
        if p.runs:
            p.runs[0].text = full_text
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                run.font.size = Pt(11)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    full_text = ''.join(run.text for run in p.runs)
                    for key, value in replace_map.items():
                        if key in full_text:
                            full_text = full_text.replace(key, value)
                    for i, run in enumerate(p.runs):
                        run.text = ''
                    if p.runs:
                        p.runs[0].text = full_text
                        for run in p.runs:
                            run.font.name = 'Times New Roman'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                            run.font.size = Pt(10)

def print_word():
    data = {label: entry.get() for label, entry in input_entries.items()}
    data.update({label: cb.get() for label, cb in comboboxes.items()})

    alo_5 = 0
    yaxshi_4 = 0
    qoniqarli_3 = 0
    qoniqarsiz_2 = 0
    kelmadi = 0

    tanlangan_guruh = data.get(translations[selected_language]["fields"]["Guruh"])
    if not tanlangan_guruh or tanlangan_guruh == translations[selected_language]["placeholder"]:
        print("❌ Guruh tanlanmagan.")
        return

    try:
        df = get_talabalar_from_google_sheet(corresponding_value)
        if df is None:
            print("❌ Talabalar ro'yxatini olishda xatolik.")
            return
        mos_talabalar = df[df['Guruh'] == tanlangan_guruh]
        sorted_talabalar = mos_talabalar[['Talaba ID', 'To‘liq ismi']].dropna(subset=['Talaba ID', 'To‘liq ismi'])
    except Exception as e:
        print("❌ Talabalarni o‘qishda xatolik:", e)
        return

    try:
        doc_name = create_word()
        doc = Document(doc_name)
    except Exception as e:
        print("❌ docx_temp.docx ochishda xatolik:", e)
        return

    nazorat_tur_full = data.get(translations[selected_language]["fields"]["Nazorat shakli"], "")
    nazorat_tur = re.sub(r"\(.*?\)", "", nazorat_tur_full).strip()

    max_ball = max_ball_map.get(nazorat_tur_full, 100)
    nazorat_shakli = data.get(translations[selected_language]["fields"]["Nazorat turi"], "")
    mapped_nazorat_shakli = nazorat_shakli_map.get(nazorat_shakli, nazorat_shakli)

    replace_map = {
        "{fakultet}": data.get(translations[selected_language]["fields"]["Fakultet nomi"], ""),
        "{semester}": data.get(translations[selected_language]["fields"]["Semestr"], ""),
        "{guruh}": data.get(translations[selected_language]["fields"]["Guruh"], ""),
        "{fan}": data.get(translations[selected_language]["fields"]["Fan"], ""),
        "{fan_uqituvchi}": data.get(translations[selected_language]["fields"]["Fan o'qituvchilari"], ""),
        "{nazorat_turi}": nazorat_tur,
        "{nazorat_tur}dan to‘plagan ballar": f"{nazorat_tur}dan to‘plagan ballar",
        "{nazorat_masuli}": data.get(translations[selected_language]["fields"]["Nazorat mas’uli"], ""),
        "{soat}": data.get(translations[selected_language]["input_fields"]["Fan soati"], ""),
        "{kredit}": data.get(translations[selected_language]["input_fields"]["Fan krediti"], ""),
        "{nazorat_sanasi}": data.get(translations[selected_language]["input_fields"]["Nazorat sanasi"], ""),
        "{tur}": mapped_nazorat_shakli,
        "{dekan}": data.get(translations[selected_language]["fields"]["Dekan nomi"], ""),
        "{mudir}": data.get(translations[selected_language]["fields"]["Kafedra mudiri nomi"], "")
    }

    for table in doc.tables:
        if any("{tr}" in cell.text for row in table.rows for cell in row.cells):
            for i, row in enumerate(table.rows):
                if "{tr}" in row.cells[0].text:
                    table._tbl.remove(row._tr)
                    break

            widths = [Cm(1), Cm(8), Cm(2.7), Cm(2.5), Cm(1.7), Cm(2)]
            for col_idx, width in enumerate(widths):
                for cell in table.columns[col_idx].cells:
                    cell.width = width

            for idx, (talaba_id, entry_widget) in enumerate(ball_entries.items(), start=1):
                parent_widgets = entry_widget.master.winfo_children()
                talaba_ismi = parent_widgets[1].cget("text") if len(parent_widgets) > 1 else ""
                ball = entry_widget.get().strip()

                ball_text = "Noto‘g‘ri"
                baho_text = "Noto‘g‘ri"

                if not ball or ball == "0":
                    ball_text = "Kelmadi"
                    baho_text = "Kelmadi"
                    kelmadi += 1
                else:
                    try:
                        ball_val = float(ball)
                        if ball_val > max_ball:
                            messagebox.showerror("Xatolik", f"❌ {talaba_ismi} uchun kiritilgan ball ({ball_val}) maksimal balldan ({max_ball}) katta.")
                            return
                        if ball_val < 0:
                            messagebox.showerror("Xatolik", f"❌ {talaba_ismi} uchun kiritilgan ball ({ball_val}) 0 dan kichik.")
                            return
                        ball_text = str(int(ball_val)) if ball_val.is_integer() else str(ball_val)
                        baho_text = calculate_baho(ball_val, max_ball)

                        if baho_text == "5":
                            alo_5 += 1
                        elif baho_text == "4":
                            yaxshi_4 += 1
                        elif baho_text == "3":
                            qoniqarli_3 += 1
                        elif baho_text == "O'tmadi":
                            qoniqarsiz_2 += 1
                        elif baho_text.lower() == "kelmadi":
                            kelmadi += 1
                    except ValueError:
                        messagebox.showerror("Xatolik", f"❌ {talaba_ismi} uchun noto‘g‘ri ball qiymati: {ball}")
                        return

                new_row = table.add_row()
                if len(new_row.cells) < 6:
                    print(f"❌ Jadval ustunlari yetarli emas. Topilgan: {len(new_row.cells)}")
                    continue

                new_row.cells[0].text = str(idx)
                new_row.cells[1].text = talaba_ismi
                new_row.cells[2].text = str(talaba_id)
                new_row.cells[3].text = ball_text
                new_row.cells[4].text = baho_text
                new_row.cells[5].text = ""

                for cell_index, cell in enumerate(new_row.cells):
                    cell.width = widths[cell_index]
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if cell_index in [3, 4] else WD_PARAGRAPH_ALIGNMENT.LEFT
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                            run.font.size = Pt(10)

            replace_map.update({
                "{student_soni}": str(len(ball_entries)),
                "{alo_5}": str(int(alo_5)),
                "{yaxshi_4}": str(int(yaxshi_4)),
                "{qoniqarli_3}": str(int(qoniqarli_3)),
                "{qoniqarsiz_2}": str(int(qoniqarsiz_2)),
                "{kelmadi}": str(int(kelmadi))
            })

            replace_text_in_doc(doc, replace_map)

    for row in doc.tables[-1].rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                    run.font.size = Pt(11)

    fan_nomi = data.get(translations[selected_language]["fields"]["Fan"], "").replace(" ", "_")
    semester = data.get(translations[selected_language]["fields"]["Semestr"], "").replace(" ", "_")
    nazorat_turi = data.get(translations[selected_language]["fields"]["Nazorat shakli"], "").replace(" ", "_")
    nazorat_shakli = data.get(translations[selected_language]["fields"]["Nazorat turi"], "").replace(" ", "_")

    output_dir = "Qaydnomalar"
    filename = f"{tanlangan_guruh}_{semester}_{fan_nomi}_{nazorat_turi}_{nazorat_shakli}_tur.docx"
    filename = re.sub(r'[\\/*?:"<>|]', "", filename)

    output_path = os.path.join(output_dir, filename)

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    doc.save(output_path)
    delete_file("docx_temp.docx")
    print(f"✅ Word hujjat yaratildi: {output_path}")
    send_file_to_telegram_group(output_path)

def send_file_to_telegram_group(file_path):
    bot_token = "7988992200:AAHguvG_iPE6ZG6gdsCjWKJ17fv6vpPzJsQ"
    chat_id = "-1002635664827"
    url = f"https://api.telegram.org/bot{bot_token}/sendDocument"

    with open(file_path, "rb") as f:
        files = {"document": f}
        data = {"chat_id": chat_id}
        response = requests.post(url, data=data, files=files)

    if response.status_code == 200:
        print("✅ Hujjat Qaydnomalar papkaga yuborildi!")
    else:
        print("❌ Hujjat Qaydnomalar papkaga yuborishda xatolik:", response.text)

def create_button(parent, text, command):
    btn = tk.Button(parent, text=text, command=command, bg=BUTTON_COLOR, fg=TEXT_COLOR, font=BUTTON_FONT,
                    bd=0, relief="flat", activebackground="#038c5a", activeforeground=BORDER_COLOR)
    btn.pack(side="left", padx=10)
    btn.configure(width=15, height=2)
    btn.bind("<Enter>", lambda e: btn.configure(bg="#038c5a"))
    btn.bind("<Leave>", lambda e: btn.configure(bg=BUTTON_COLOR))
    return btn

# Show language selection dialog
select_language()

root.mainloop()