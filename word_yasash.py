from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

def create_word():
   
    # Yangi Word hujjati yaratish
    doc = Document()

    # Sahifa chetlarini sozlash (1 inch = 2.54 sm)
    # Sahifa chetlarini sozlash
    section = doc.sections[0]
    section.left_margin = Cm(2)    # Chapdan 2 sm
    section.right_margin = Cm(1.5) # O‘ngdan 1.5 sm
    section.top_margin = Cm(2)     # Yuqoridan 2 sm
    section.bottom_margin = Cm(2)  # Pastdan 2 sm

    # Hujjat sarlavhasi
    # Hujjat sarlavhasi
    title = doc.add_paragraph()
    title_run = title.add_run("{tur}-shakl")
    title_run.bold = False
    title_run.font.name = 'Times New Roman'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    title_run.font.size = Pt(11)
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # ← ONG TOMONGA
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(0)
    title.paragraph_format.line_spacing = 1.0

    # Universitet nomi
    uni = doc.add_paragraph()
    uni_run = uni.add_run("OSIYO XALQARO UNIVERSITETI")
    uni_run.bold = True
    uni_run.font.name = 'Times New Roman'
    uni_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    uni_run.font.size = Pt(12)
    uni.alignment = WD_ALIGN_PARAGRAPH.CENTER
    uni.paragraph_format.space_before = Pt(0)
    uni.paragraph_format.space_after = Pt(0)
    uni.paragraph_format.line_spacing = 1.0

    # Baholash qaydnomasi
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("BAHOLASH QAYDNOMASI № ______\n")
    subtitle_run.bold = True
    subtitle_run.font.name = 'Times New Roman'
    subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    subtitle_run.font.size = Pt(12)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(0)
    subtitle.paragraph_format.line_spacing = 1.0

    # Fakultet, semestr va guruh ma'lumotlari
    info = doc.add_paragraph()
    info_run = info.add_run("Fakultet: {fakultet}, Semestr: {semester}, Guruh: {guruh}")
    info_run.font.name = 'Times New Roman'
    info_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    info_run.font.size = Pt(11)
    info.alignment = WD_ALIGN_PARAGRAPH.LEFT
    info.paragraph_format.space_before = Pt(0)
    info.paragraph_format.space_after = Pt(0)
    info.paragraph_format.line_spacing = 1.5

    # Fan va o'qituvchi ma'lumotlari
    fan = doc.add_paragraph()
    fan_run = fan.add_run("Fan: {fan}")
    fan_run.font.name = 'Times New Roman'
    fan_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    fan_run.font.size = Pt(11)
    fan.paragraph_format.space_before = Pt(0)
    fan.paragraph_format.space_after = Pt(0)
    fan.paragraph_format.line_spacing = 1.5

    teacher = doc.add_paragraph()
    teacher_run = teacher.add_run("Fan o‘qituvchilari: {fan_uqituvchi}")
    teacher_run.font.name = 'Times New Roman'
    teacher_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    teacher_run.font.size = Pt(11)
    teacher.paragraph_format.space_before = Pt(0)
    teacher.paragraph_format.space_after = Pt(0)
    teacher.paragraph_format.line_spacing = 1.5

    control = doc.add_paragraph()
    control_run = control.add_run("{nazorat_turi} nazorat mas’uli: {nazorat_masuli}")
    control_run.font.name = 'Times New Roman'
    control_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    control_run.font.size = Pt(11)
    control.paragraph_format.space_before = Pt(0)
    control.paragraph_format.space_after = Pt(0)
    control.paragraph_format.line_spacing = 1.5

    hours = doc.add_paragraph()
    hours_run = hours.add_run("Semestrda fanga ajratilgan umumiy soatlar/kredit: {soat} / {kredit}")
    hours_run.font.name = 'Times New Roman'
    hours_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    hours_run.font.size = Pt(11)
    hours.paragraph_format.space_before = Pt(0)
    hours.paragraph_format.space_after = Pt(0)
    hours.paragraph_format.line_spacing = 1.5

    date = doc.add_paragraph()
    date_run = date.add_run("{nazorat_turi} nazorat/qayta topshirish o‘tkazilgan sana: {nazorat_sanasi}")
    date_run.font.name = 'Times New Roman'
    date_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    date_run.font.size = Pt(11)
    date.paragraph_format.space_before = Pt(0)
    date.paragraph_format.space_after = Pt(0)
    date.paragraph_format.line_spacing = 1.0

    # 1 ta Enter (bo‘sh paragraf) qo‘shish
    doc.add_paragraph("")

    # Jadval yaratish
    # Jadval yaratish
    table = doc.add_table(rows=2, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    table.autofit = False  # MUHIM! AutoFit'ni o‘chirib qo‘yamiz

    # Har bir ustun uchun kenglik (aniq)
    widths = [Cm(1), Cm(6), Cm(3), Cm(3), Cm(2), Cm(3)]
    for col, width in zip(table.columns, widths):
        for cell in col.cells:
            cell.width = width

    # 1-qator: sarlavhalar
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "\n№"
    hdr_cells[1].text = "\nTalabaning \nfamiliyasi, \nismi, sharifi"
    hdr_cells[2].text = "\nReyting daftar-chasining raqami"
    hdr_cells[3].text = "\n{nazorat_tur}dan to‘plagan ballar"
    hdr_cells[4].text = "\nBaho"
    hdr_cells[5].text = "\n{nazorat_turi}\no‘tkazgan o‘qituvchi imzosi"

    # 2-qator: misol uchun ma'lumot
    row_cells = table.rows[1].cells
    row_cells[0].text = "{tr}"
    row_cells[1].text = "{talaba_nomi}"
    row_cells[2].text = "{hemis_id}"
    row_cells[3].text = "{Ball}"
    row_cells[4].text = "{Baho}"
    row_cells[5].text = ""

    # Har bir katakdagi matnni formatlash
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

                    run.font.size = Pt(10)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0  # 1.5 dan 1.0 ga o‘zgartirildi

    # Jadvaldan keyin 2 ta Enter (bo‘sh paragraf) qo‘shish
    doc.add_paragraph("")
    # doc.add_paragraph("")

    # Jami talabalar statistikasi
    stats = doc.add_paragraph()
    stats_run = stats.add_run(
        "Jami talabalar soni: {student_soni}, shundan: “5”: {alo_5}, “4”: {yaxshi_4}, "
        "“3”: {qoniqarli_3}, “O'tmadi”: {qoniqarsiz_2}, Kelmadi: {kelmadi}\n"
    )
    stats_run.font.name = 'Times New Roman'
    stats_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    stats_run.font.size = Pt(11)
    stats.paragraph_format.space_before = Pt(0)
    stats.paragraph_format.space_after = Pt(0)
    stats.paragraph_format.line_spacing = 1.0

    # Imzolar uchun jadval
    sign_table = doc.add_table(rows=2, cols=3)
    sign_table.autofit = False
    sign_table.allow_autofit = False
    sign_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Ustun kengliklarini qat'iy belgilaymiz
    sign_table.cell(0, 0).width = Cm(3)
    sign_table.cell(0, 1).width = Cm(3)
    sign_table.cell(0, 2).width = Cm(10)
    sign_table.cell(1, 0).width = Cm(3)
    sign_table.cell(1, 1).width = Cm(3)
    sign_table.cell(1, 2).width = Cm(10)

    # Kataklardagi matn
    sign_table.cell(0, 0).text = "Fakultet dekani"
    sign_table.cell(0, 1).text = "_____________"
    sign_table.cell(0, 2).text = "{dekan}\n"

    sign_table.cell(1, 0).text = "Kafedra mudiri"
    sign_table.cell(1, 1).text = "_____________"
    sign_table.cell(1, 2).text = "{mudir}"

    # Formatlash: shrift va joylashuv
    for row in sign_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Hujjatni saqlash
    doc.save("docx_temp.docx")
    return "docx_temp.docx"

def delete_file(file_name):
    if os.path.exists(file_name):
        os.remove(file_name)
        print("O'chirildi")
    else:
        print("file topilmadi")